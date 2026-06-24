"""
Generate a Word (.docx) version of the Analytics Activity Ingest API reference
for frontend developers and QA testers.

Run: python scripts/gen_activity_api_doc.py
Output: Pathzi_Analytics_Activity_API.docx (project root)
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0F, 0x2E, 0x4F)
BLUE = RGBColor(0x1F, 0x6F, 0xB2)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5)


def h1(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY


def h2(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BLUE


def body(t, italic=False, color=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(t); r.italic = italic; r.font.size = Pt(10.5)
    if color: r.font.color.rgb = color


def bullet(t):
    doc.add_paragraph(t, style="List Bullet").paragraph_format.space_after = Pt(2)


def code(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(head); run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(255, 255, 255)
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "1F6FB2"); c._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(8.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ---- cover ----
doc.add_paragraph().paragraph_format.space_before = Pt(40)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Pathzi Analytics"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = NAVY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Activity Ingest API — Reference"); r.font.size = Pt(16); r.font.color.rgb = BLUE
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run("For Frontend Developers & QA Testers"); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY
doc.add_page_break()

# 1 auth
h1("1. Authentication")
body("Both endpoints require an authenticated user. Send the JWT access token in the header:")
code("Authorization: Bearer <access_token>\nContent-Type: application/json")
bullet("The user is taken from the token — never send a user id in the body (it is ignored).")
bullet("Unauthenticated requests return 401 Unauthorized.")

# 2 activity
h1("2. POST /analytics/activity/  — report events")
body("Send a batch of 1–200 events. The server validates, queues, and returns immediately. "
     "Events are written asynchronously and appear in reports within ~30 seconds.")

h2("2.1 Request body")
code('{\n  "events": [\n'
     '    {"activity_type":"career_viewed","career_id":2336,\n'
     '     "metadata":{"source":"swipe_deck","position":4,"session_id":"abc123"}},\n'
     '    {"activity_type":"career_swiped_right","career_id":2336,\n'
     '     "metadata":{"source":"swipe_deck","position":4}}\n  ]\n}')

h2("2.2 Event fields")
table(["Field", "Type", "Required", "Notes"],
      [["activity_type", "string", "Yes", "One of the allowed values in 2.3"],
       ["career_id", "integer", "No", "DB id of the career the action is about"],
       ["route_id", "string", "No", "Education-route type (route events): course | apprenticeship | job"],
       ["activity_value", "string", "No", "Optional label (provider/route name)"],
       ["card", "string", "No", "Title of the course/apprenticeship/job card (provider clicks)"],
       ["metadata", "object", "No", "Free-form context. Max 50 keys. No personal data"]],
      widths=[1.3, 0.9, 0.8, 3.5])

h2("2.3 Allowed activity_type values (frontend-fired only)")
table(["activity_type", "When the app sends it", "Recommended fields"],
      [["career_viewed", "A career card is shown", "career_id; meta: source, position, session_id"],
       ["career_swiped_right", "User likes / swipes right", "career_id; meta: source, position"],
       ["career_swiped_left", "User skips / swipes left", "career_id; meta: source, position"],
       ["route_viewed", "An education route is shown", "career_id, route_id=course|apprenticeship|job"],
       ["route_clicked", "User taps an education route", "career_id, route_id=course|apprenticeship|job"],
       ["provider_link_clicked", "User opens a provider link", "career_id, route_id, card=title; activity_value=provider"],
       ["connect_button_clicked", "User taps a connect button", "career_id, route_id, card=title; activity_value=provider"]],
      widths=[1.7, 2.2, 2.6])
body("Do NOT send these (recorded automatically by the backend; rejected with 400): "
     "career_saved, career_unsaved, career_explored, career_unexplored, search_performed. "
     "consent_given has its own endpoint (section 3).", italic=True, color=GREY)

h2("2.4 Success response — 202 Accepted")
code('{ "status": true, "queued": 2 }')
body("queued = number of events accepted into the processing queue.")

h2("2.5 Limits")
table(["Limit", "Value"],
      [["Events per request", "1–200"], ["Metadata keys per event", "50"],
       ["Rate limit", "120 requests / minute / user"]], widths=[3.0, 3.5])

h2("2.6 curl example")
code('curl -X POST https://pathzi.co.uk/analytics/activity/ \\\n'
     '  -H "Authorization: Bearer <access_token>" \\\n'
     '  -H "Content-Type: application/json" \\\n'
     '  -d \'{"events":[{"activity_type":"career_viewed","career_id":2336,\n'
     '       "metadata":{"source":"swipe_deck","position":4}}]}\'')

h2("2.7 Flutter / Dart example")
code('final res = await http.post(\n'
     "  Uri.parse('https://pathzi.co.uk/analytics/activity/'),\n"
     "  headers: {'Authorization': 'Bearer \$accessToken',\n"
     "           'Content-Type': 'application/json'},\n"
     "  body: jsonEncode({'events': [\n"
     "    {'activity_type': 'career_viewed', 'career_id': 2336,\n"
     "     'metadata': {'source': 'swipe_deck', 'position': 4}}\n"
     "  ]}),\n);")

h2("2.8 Recommended client pattern (batching)")
bullet("Buffer events on the device; flush in batches, not one request per event.")
bullet("Flush at ~20–50 events, on app background, or every few seconds.")
bullet("Keep each batch ≤ 200 events.")
bullet("On failure (offline), keep events and retry on next launch.")

# 3 consent
h1("3. POST /analytics/consent/  — record consent")
body("Records explicit consent for a provider to contact the user. This is the ONLY endpoint that "
     "accepts contact data. Written synchronously; also logs an anonymous consent_given event (no email).")
h2("3.1 Request body")
code('{\n  "career_id": 2336,\n  "provider_name": "City College",\n'
     '  "provider_type": "college",\n  "contact_email": "user@example.com"\n}')
table(["Field", "Type", "Required", "Notes"],
      [["career_id", "integer", "No", "Must exist if provided"],
       ["provider_name", "string", "Yes", "Max 255 chars"],
       ["provider_type", "string", "No", "college, employer, training_provider, ..."],
       ["contact_email", "string", "Yes", "Valid email"]], widths=[1.3, 0.9, 0.8, 3.5])
h2("3.2 Success response — 201 Created")
code('{\n  "status": true,\n  "message": "Consent recorded.",\n'
     '  "lead_id": 17,\n  "consent_at": "2026-06-16T13:40:00Z"\n}')

# 4 errors
h1("4. Error responses")
table(["Status", "Meaning", "Example"],
      [["400", "Validation failed (bad/empty events, disallowed type, >200, invalid email)", "{events: {...}} / field errors"],
       ["401", "Missing/invalid token", '{"detail":"Authentication credentials were not provided."}'],
       ["429", "Rate limit exceeded (>120/min)", '{"detail":"Request was throttled..."}']],
      widths=[0.8, 3.0, 2.7])
h2("Common 400 cases (activity endpoint)")
bullet("events missing or empty → 'events must contain at least 1 event.'")
bullet("More than 200 events → 'Too many events …'")
bullet("Unknown / backend-only activity_type → 'Unsupported activity_type …'")
bullet("metadata with >50 keys → 'metadata has too many keys (max 50).'")
body("An invalid / non-existent career_id does NOT fail the request — the event is stored and the "
     "broken career link is dropped (null). This keeps high-volume ingest resilient.", italic=True, color=GREY)

# 5 privacy
h1("5. Data & Privacy Rules (GDPR)")
bullet("Never put personal data (emails, phones, names) in metadata.")
bullet("The server auto-redacts emails/phones in metadata to [redacted] as a safety net — don't rely on it.")
bullet("Contact data belongs only in /analytics/consent/, only after explicit consent.")
bullet("Views, clicks, swipes, routes are stored as anonymous analytics.")

# 6 behaviour
h1("6. Behaviour Notes (for testers)")
bullet("Asynchronous: /activity/ returns 202 instantly; events appear within ~30 seconds (background flush every 30s).")
bullet("Synchronous: /consent/ writes immediately and returns the new lead_id.")
bullet("Dependency: queued activity events require the backend worker to be running; consent does not.")
bullet("Identity: events are always attributed to the authenticated user from the token.")

# 7 QA checklist
h1("7. QA Test Checklist")
h2("/analytics/activity/")
for t in ["Valid single-event batch → 202, queued:1; row appears within 30s",
          "Valid 200-event batch → 202, queued:200",
          "201 events → 400 (too many)",
          "Empty events: [] → 400",
          "Missing events key → 400",
          "Disallowed type (career_saved) → 400",
          "Unknown type (foo_bar) → 400",
          "Email in metadata → stored with email redacted",
          "Non-existent career_id → still 202; stored without career link",
          "No auth header → 401",
          ">120 requests in a minute → 429"]:
    bullet(t)
h2("/analytics/consent/")
for t in ["Valid body → 201 with lead_id + consent_at; ProviderLead created",
          "Missing contact_email → 400",
          "Invalid email → 400",
          "Non-existent career_id → 400",
          "Consent event recorded in analytics WITHOUT the email",
          "No auth header → 401"]:
    bullet(t)

# 8 quick ref
h1("8. Quick Reference")
code("POST /analytics/activity/  Auth: Bearer  Body:{events:[...]}  -> 202 {queued:N}\n"
     "POST /analytics/consent/   Auth: Bearer  Body:{provider_name,contact_email,...} -> 201 {lead_id}\n\n"
     "Allowed activity_type (activity endpoint):\n"
     "  career_viewed, career_swiped_right, career_swiped_left,\n"
     "  route_viewed, route_clicked, provider_link_clicked, connect_button_clicked\n\n"
     "Limits: 1-200 events/request | 50 metadata keys/event | 120 requests/min\n"
     "Events appear in reports within ~30 seconds.")

out = "Pathzi_Analytics_Activity_API.docx"
doc.save(out)
print("WROTE", out)
