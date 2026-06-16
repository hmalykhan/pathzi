"""
Generates the client-facing Word document describing Pathzi's
Retrieval-Augmented (RAG-based) Career Recommendation Engine, including the
dedicated embedding microservice deployed on a DigitalOcean droplet.

Run: python scripts/gen_rag_doc.py
Output: Pathzi_Career_Recommendation_RAG.docx (project root)
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0F, 0x2E, 0x4F)
BLUE = RGBColor(0x1F, 0x6F, 0xB2)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x8A, 0x8A, 0x8A)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def _heading(text, size, color, space_before=14, space_after=6, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    return p


def h1(text):
    return _heading(text, 16, NAVY)


def h2(text):
    return _heading(text, 13, BLUE)


def body(text, italic=False, color=None, size=11, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F6FB2")
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def divider():
    p = doc.add_paragraph()
    r = p.add_run("―" * 40)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============================ COVER ============================
doc.add_paragraph().paragraph_format.space_before = Pt(90)
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Pathzi"); r.font.size = Pt(40); r.font.bold = True; r.font.color.rgb = NAVY
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("AI-Powered Career Recommendation Engine"); r.font.size = Pt(20); r.font.color.rgb = BLUE
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Retrieval-Augmented (RAG-based) Architecture"); r.font.size = Pt(14); r.italic = True; r.font.color.rgb = GREY
doc.add_paragraph(); doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Technical Overview & Architecture Document\n"
                 "Embeddings · Vector Search · Dedicated ML Microservice\n\nPrepared for the Client")
r.font.size = Pt(12); r.font.color.rgb = GREY
doc.add_page_break()

# ============================ 1. EXEC SUMMARY ============================
h1("1. Executive Summary")
body(
    "Pathzi helps young people discover careers that genuinely fit them. At the heart of the platform "
    "is an AI-powered recommendation engine that learns who each user is — their education, interests, "
    "and in-app behaviour — and matches them to the most relevant careers from Pathzi's career library."
)
body(
    "The engine uses a Retrieval-Augmented (RAG-based) architecture. Both the user and every career are "
    "converted into mathematical 'meaning vectors' (embeddings), and the system retrieves the careers "
    "whose meaning is closest to the user. The platform runs this with two cooperating parts:"
)
bullet(" the main Pathzi backend, which builds the text profiles, stores embeddings, runs the vector "
       "search, and serves recommendations to the app.", "Application backend — ")
bullet(" a dedicated AI embedding microservice deployed on a DigitalOcean droplet, which hosts the "
       "language model and converts text into embeddings on demand.", "Embedding microservice — ")
body(
    "Crucially, the entire career library was embedded in advance and stored in the database, so that "
    "live recommendations are a fast vector search rather than a slow re-computation. This is the same "
    "design pattern used by modern semantic-search and retrieval systems.",
    after=4,
)

# ============================ 2. WHAT RAG MEANS ============================
h1("2. What 'Retrieval-Augmented' Means Here")
body("A retrieval-augmented system works in two phases — an offline phase and an online phase:")
bullet(" every career in the library is turned into an embedding once and stored in a vector "
       "database. This is the searchable 'knowledge index'.", "Offline (build the index) — ")
bullet(" when a user needs recommendations, their profile is turned into an embedding and the system "
       "retrieves the closest careers from that index.", "Online (answer a query) — ")
body(
    "Because matches are based on meaning rather than exact keywords, a user interested in 'helping "
    "people and healthcare' can be matched to relevant roles even if their profile never used those "
    "precise words — the embeddings place similar concepts near each other.",
    after=4,
)
body(
    "Scope note: the engine currently implements the retrieval and intelligent-ranking stages of the "
    "RAG paradigm (embeddings + semantic vector search + diversification). It does not yet generate "
    "free-text narratives with a language model; the architecture is ready to be extended with a "
    "generative layer in future if desired.",
    italic=True, color=GREY, size=10,
)

# ============================ 3. SYSTEM COMPONENTS ============================
h1("3. System Components")
table(
    ["Component", "Where it runs", "Responsibility"],
    [
        ["Application backend", "Pathzi Django server",
         "Builds user/career text, stores embeddings, runs vector search, serves results"],
        ["Embedding microservice", "Dedicated DigitalOcean droplet",
         "Hosts the AI model; converts text into 384-dim embeddings via an HTTP endpoint"],
        ["Vector database", "PostgreSQL + pgvector",
         "Stores all user and career embeddings; performs fast similarity search"],
    ],
    widths=[1.9, 2.0, 2.6],
)
body(
    "Separating the embedding model into its own microservice keeps the AI workload isolated from the "
    "main application, lets it be scaled or upgraded independently, and guarantees that users and "
    "careers are always embedded by the exact same model — which is essential for accurate matching.",
    after=4,
)

# ============================ 4. EMBEDDING MICROSERVICE ============================
h1("4. The Embedding Microservice (DigitalOcean Droplet)")
body(
    "The embedding microservice is a small, focused Django service whose only job is to turn text into "
    "an embedding. It exposes a single HTTP endpoint that the main backend calls whenever it needs to "
    "embed a user profile or a career."
)
h2("4.1 The AI model is loaded once")
body(
    "The service loads the language model a single time and keeps it in memory (a 'singleton'). Loading "
    "such a model is expensive, so loading it once and reusing it makes every subsequent request fast "
    "and efficient."
)
h2("4.2 The endpoint")
table(
    ["Property", "Value"],
    [
        ["Method & path", "POST /ml/embed/"],
        ["Request body", '{ "text": "<the text to embed>" }'],
        ["Successful response", '{ "embedding": [ 384 floating-point numbers ] }'],
        ["Missing text", "HTTP 400 — 'Text is required'"],
        ["Internal error", "HTTP 500 — error message"],
        ["Model", "sentence-transformers/all-MiniLM-L6-v2 (384 dimensions)"],
    ],
    widths=[2.0, 4.5],
)
body("Example exchange:", after=2)
code_block(
    "REQUEST   POST /ml/embed/\n"
    '          { \"text\": \"Education level A-Levels; interest in Healthcare; ...\" }\n\n'
    "RESPONSE  200 OK\n"
    '          { \"embedding\": [0.0123, -0.0457, 0.0991, ... ] }   (384 numbers)'
)
body(
    "The service also converts the model's raw output into a plain list of numbers so it can be stored "
    "directly in the pgvector database.",
    after=4,
)

# ============================ 5. OFFLINE PHASE ============================
h1("5. Offline Phase — Embedding the Entire Career Library")
body(
    "Before any recommendation can be made, Pathzi builds its searchable knowledge index. Every career "
    "in the library is processed once and stored as an embedding:"
)
numbered(" each career is converted into a clean, structured text summary (title, category, "
         "description, how-to-become, education and apprenticeship routes, salary and hours).",
         "Describe the career — ")
numbered(" that text is sent to the embedding microservice and converted into a 384-dimension vector.",
         "Embed it — ")
numbered(" the embedding is saved next to the career in the pgvector database (the CareerEmbedding "
         "store).", "Store it — ")
body(
    "Because this is done in advance, the heavy AI work is already complete before any user asks for "
    "recommendations. At request time the system only has to embed the one user and compare — which is "
    "extremely fast. This 'index once, query many times' design is what makes the experience responsive.",
    after=4,
)

# ============================ 6. ONLINE PHASE ============================
h1("6. Online Phase — Generating a Recommendation")
body("When a user needs recommendations, the engine runs five quick stages:")
numbered(" gather the user's profile (education level, age, interest categories) and behaviour "
         "(careers they saved or explored).", "Understand the user — ")
numbered(" write this into a clean paragraph describing the user's career intent.",
         "Build a meaning profile — ")
numbered(" send that text to the embedding microservice to get the user's 384-dim embedding, and store "
         "it.", "Create the user embedding — ")
numbered(" compare the user's embedding against the pre-built career embeddings and retrieve the "
         "closest matches using vector search.", "Retrieve the best matches — ")
numbered(" remove near-duplicate results, cache the top list, and serve it to the app instantly.",
         "Refine, cache, and serve — ")
body(
    "Whenever the user saves or explores a career, this loop quietly re-runs in the background so "
    "recommendations stay current — without ever slowing down the app.",
    after=4,
)

h2("6.1 Architecture at a Glance")
code_block(
    "                    EMBEDDING MICROSERVICE  (DigitalOcean droplet)\n"
    "                    POST /ml/embed/  ->  all-MiniLM-L6-v2  ->  384-dim vector\n"
    "                         ^                              ^\n"
    "         (offline) career text                 (online) user text\n"
    "                         |                              |\n"
    "  CAREER LIBRARY --------+                              +-------- USER profile + saved/explored\n"
    "         |                                                            |\n"
    "         v                                                            v\n"
    "  CareerEmbedding (pgvector)  <==== VECTOR SIMILARITY SEARCH ====>  UserEmbedding (pgvector)\n"
    "                                        (cosine distance)\n"
    "                                              |\n"
    "                                              v\n"
    "                                  Diversify (drop near-duplicates)\n"
    "                                              |\n"
    "                                              v\n"
    "                                  Cache top 50  ->  Serve to mobile app"
)

# ============================ 7. USER REPRESENTATION ============================
h1("7. Building the User's Meaning Profile")
body(
    "Pathzi converts each user into structured text that captures their career intent, combining three "
    "layers of signal:"
)
table(
    ["Signal layer", "Source", "Strength"],
    [
        ["Profile attributes", "Education level, age/stage, interest categories", "Baseline"],
        ["Saved careers", "Careers the user explicitly saved", "Strong signal"],
        ["Explored careers", "Careers the user opened/looked at", "Weaker signal"],
    ],
    widths=[2.2, 3.0, 1.4],
)
body("Saved careers count as stronger evidence than explored ones, and junk/duplicate values are "
     "removed. The resulting profile looks like this:")
code_block(
    "Career recommendation profile\n\n"
    "Profile:\nEducation level: A-Levels.\nAge group or stage: 17.\n\n"
    "Interest categories:\n- Healthcare\n- Helping people\n\n"
    "Strong signals from saved careers:\n- Paramedic\n- Nurse\n\n"
    "Overall career intent:\nThis user has education level A-Levels; interest in Healthcare, "
    "Helping people; strong interest in Paramedic, Nurse."
)

# ============================ 8. CAREER REPRESENTATION ============================
h1("8. Building Each Career's Meaning Profile")
body("Every career is converted into structured text, with the most important signals first:")
bullet(" the career title (strongest signal)", "Title — ")
bullet(" sector and sub-type", "Category — ")
bullet(" what the role involves", "Description — ")
bullet(" the realistic entry routes", "How to become — ")
bullet(" college courses and entry requirements", "Education path — ")
bullet(" apprenticeship routes and requirements", "Apprenticeship path — ")
bullet(" salary, hours, and working pattern", "Work details — ")
body("This text is embedded (offline) and stored with the career, ready for instant comparison.", after=4)

# ============================ 9. MODEL ============================
h1("9. The AI Embedding Model")
table(
    ["Property", "Value"],
    [
        ["Model", "sentence-transformers/all-MiniLM-L6-v2"],
        ["Embedding size", "384 dimensions"],
        ["Served by", "Dedicated microservice on a DigitalOcean droplet"],
        ["Similarity measure", "Cosine distance"],
        ["Vector storage", "PostgreSQL with the pgvector extension"],
    ],
    widths=[2.2, 4.3],
)
body(
    "Users and careers are embedded by the same model via the same microservice, placing them in one "
    "shared 'meaning space' so they can be compared directly. The model is compact and fast, keeping "
    "the service responsive and cost-efficient.",
    after=4,
)

# ============================ 10. RETRIEVAL ============================
h1("10. Retrieval & Diversification")
body("With the user represented as an embedding, the engine searches the pre-built career index:")
bullet(" careers are first narrowed to the user's chosen interest categories, keeping results "
       "relevant.", "Category-aware — ")
bullet(" careers are ranked by cosine distance — smaller distance means closer meaning.",
       "Semantic ranking — ")
bullet(" a wider candidate pool (five times the final size) is retrieved so there is room to refine "
       "for quality.", "Over-fetch for quality — ")
h2("10.1 Diversification")
body(
    "Pure similarity search can surface many near-identical careers. To keep the list useful and "
    "varied, the engine removes any candidate that is more than 95% similar to one already chosen, "
    "then trims to the top results (currently up to 50)."
)

# ============================ 11. FRESHNESS ============================
h1("11. Staying Fresh — The Self-Updating Loop")
bullet(" saving or exploring a career schedules a background refresh of the user's embedding and "
       "recommendations.", "Triggered by behaviour — ")
bullet(" rapid repeated actions collapse into a single refresh, saving processing and cost.",
       "Debounced — ")
bullet(" heavy work runs in the background, so the user's action returns instantly.", "Non-blocking — ")
bullet(" the final list is cached for fast delivery (refreshed every 6 hours, or sooner when active).",
       "Cached — ")

# ============================ 12. BENEFITS ============================
h1("12. Why This Approach Benefits Pathzi")
table(
    ["Capability", "Business benefit"],
    [
        ["Understands meaning, not keywords", "Better, more human-feeling matches"],
        ["Pre-built career index", "Fast recommendations, low latency"],
        ["Dedicated embedding service", "AI scales independently of the app"],
        ["Learns from user behaviour", "Recommendations improve with engagement"],
        ["Diversified results", "Users discover a broader range of careers"],
        ["Background processing & caching", "App stays fast under load"],
        ["Extensible architecture", "Ready for future AI enhancements"],
    ],
    widths=[2.8, 3.7],
)

# ============================ 13. FUTURE ============================
h1("13. Future Enhancements (Optional)")
bullet(" add a generative layer that writes a short, personalised explanation of why each career was "
       "recommended (full RAG with generation).", "Explanations — ")
bullet(" factor location into matching so nearby opportunities rank higher.", "Location-aware matching — ")
bullet(" use new analytics signals (views, swipes, clicks) to sharpen personalisation.",
       "Behavioural depth — ")
bullet(" periodically measure recommendation quality against real engagement.",
       "Continuous evaluation — ")

# ============================ 14. GLOSSARY ============================
h1("14. Glossary")
table(
    ["Term", "Plain-language meaning"],
    [
        ["Embedding", "A list of numbers that captures the meaning of text so a computer can compare it."],
        ["Microservice", "A small, independent service that does one job — here, creating embeddings."],
        ["Vector search", "Finding items whose meaning-numbers are closest to each other."],
        ["Cosine distance", "A measure of how close two embeddings are in meaning."],
        ["pgvector", "A PostgreSQL extension that stores and searches embeddings efficiently."],
        ["Offline / online", "Offline = build the index in advance; online = answer a user query live."],
        ["RAG", "Retrieval-Augmented approach: use embeddings to retrieve the most relevant items."],
        ["Diversification", "Removing near-duplicate results so the list is varied."],
    ],
    widths=[1.7, 4.8],
)

divider()
foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Pathzi — AI Career Recommendation Engine · Technical Overview")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = LIGHT

out = "Pathzi_Career_Recommendation_RAG.docx"
doc.save(out)
print("WROTE", out)
